"""
앨리영어 - 1인 학원 맞춤형 AI 종합 관리 대시보드 시스템
Streamlit 웹 애플리케이션
"""
import streamlit as st
import pandas as pd
import json
import os
import io
import base64
import calendar
from datetime import datetime, timedelta
from PIL import Image
from docx import Document
import altair as alt
from dotenv import load_dotenv

# .env 파일 로드
load_dotenv()

# 내부 모듈
from db_manager import *
from ai_service import *
from ai_service import _model_cache
from messaging import send_solapi_message, send_message, send_kakao_alimtalk, check_solapi_balance, send_comprehensive_alimtalk

# ─── 페이지 설정 ─────────────────────────────────────────────
st.set_page_config(
    page_title="대치앨리영어 AI 대시보드",
    page_icon="🎀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── 스타일링 ──────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;700&display=swap');
    html, body, [data-testid="stSidebar"] { font-family: 'Noto Sans KR', sans-serif; }
    .stApp { background-color: white !important; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #D81B60; padding: 0.5rem 0; }
    .sub-header { font-size: 1.5rem; font-weight: 600; color: #AD1457; padding: 0.3rem 0; border-bottom: 2px solid #FCE4EC; margin-bottom: 1rem; }
    .stButton>button { border-radius: 20px; transition: all 0.3s; font-weight: 600; width: 100%; }
    [data-testid="stSidebar"] { background-color: #FFF5F8 !important; border-right: 1px solid #FCE4EC; }
    .info-card { background: white; padding: 1.2rem; border-radius: 15px; box-shadow: 0 2px 10px rgba(0,0,0,0.05); border: 1px solid #FCE4EC; margin-bottom: 1rem; }
    .stat-card { background: linear-gradient(135deg, #FFF5F8, #FCE4EC); padding: 1.2rem; border-radius: 15px; text-align: center; border: 1px solid #F8BBD0; }
    .stat-number { font-size: 2.2rem; font-weight: 800; color: #D81B60; }
    .stat-label { font-size: 0.85rem; color: #888; }
    .calendar-day { padding: 8px; border: 1px solid #f0f0f0; text-align: center; min-height: 80px; font-size: 0.8rem; }
    .calendar-day.present { background: #E8F5E9; border-color: #A5D6A7; }
    .calendar-day.absent { background: #FFEBEE; border-color: #EF9A9A; }
    .calendar-day.late { background: #FFF3E0; border-color: #FFCC80; }
    .calendar-day.video { background: #E3F2FD; border-color: #90CAF9; }
    .report-card { background: white; padding: 2rem; border-radius: 20px; border: 2px solid #FCE4EC; box-shadow: 0 4px 15px rgba(0,0,0,0.05); max-width: 800px; margin: 0 auto; }
</style>
""", unsafe_allow_html=True)

# ─── 세션 상태 초기화 ──────────────────────────────────────
if "db_initialized" not in st.session_state:
    init_database()
    st.session_state.db_initialized = True
if "menu" not in st.session_state:
    st.session_state.menu = "📊 대시보드"
if "last_exam_result" not in st.session_state:
    st.session_state.last_exam_result = None
if "vocab_result" not in st.session_state:
    st.session_state.vocab_result = None
if "sel_date" not in st.session_state:
    st.session_state.sel_date = datetime.now().strftime("%Y-%m-%d")
if "monthly_summary" not in st.session_state:
    st.session_state.monthly_summary = None
if "monthly_avg" not in st.session_state:
    st.session_state.monthly_avg = None

# ─── 상단 로고 & 헤더 ───────────────────────────────
st.markdown(f"""
<div style="display:flex; justify-content:space-between; align-items:center; padding:1rem; background:white; border-radius:15px; margin-bottom:1.5rem; border:1px solid #FCE4EC;">
    <div style="display:flex; align-items:center; gap:15px;">
        <span style="font-size:2.5rem;">🎀</span>
        <div>
            <div style="font-size:1.8rem; font-weight:800; color:#D81B60; line-height:1.2;">대치앨리영어</div>
            <div style="font-size:0.9rem; color:#888;">For My Beloved Allie</div>
        </div>
    </div>
    <div style="text-align:right;">
        <div style="font-size:0.9rem; color:#666; font-weight:500;">{datetime.now().strftime('%Y년 %m월 %d일')}</div>
        <div style="font-size:0.75rem; color:#AD1457;">오늘도 아이들과 행복한 시간 되세요 원장님! ✨</div>
    </div>
</div>
""", unsafe_allow_html=True)

# ─── 사이드바 네비게이션 ───────────────────────────────
with st.sidebar:
    st.markdown("<div style='text-align:center; padding:0.5rem;'><span style='font-size:4rem;'>👩‍🏫</span></div>", unsafe_allow_html=True)
    
    nav_menus = [
        "📊 대시보드",
        "📅 출석부",
        "⚙️ 학원 관리",
         "📖 교재분석",
        "🔬 시험지 OCR",
        "📈 성적 & 리포트",
        "💰 횟수 관리",
        "🏫 학교별 기출",
        "📋 히스토리",
        "📨 알림톡 발송"
    ]
    for m in nav_menus:
        if st.button(m, use_container_width=True, type="primary" if st.session_state.menu == m else "secondary"):
            st.session_state.menu = m
            st.rerun()

    st.divider()
    with st.expander("⚙️ 시스템 설정"):
        env_key = os.getenv("GEMINI_API_KEY", "")
        st.text_input("Gemini API Key", value=env_key, type="password", key="gemini_key_input")
        st.text_input("Solapi Key", value=os.getenv("SOLAPI_API_KEY", ""), key="solapi_key_input")
        st.text_input("Solapi Secret", value=os.getenv("SOLAPI_API_SECRET", ""), type="password", key="solapi_secret_input")
        st.text_input("발신 번호", value=os.getenv("SOLAPI_FROM_NUMBER", ""), key="solapi_from_input")
        # API 키 상태 표시
        current_key = st.session_state.get("gemini_key_input", "")
        if current_key:
            st.success(f"✅ Gemini API 키 설정됨 (길이: {len(current_key)}자)")
            # 현재 사용 중인 모델 표시
            cached_model = _model_cache.get(current_key, "자동 선택")
            st.info(f"🤖 현재 모델: {cached_model}")
        else:
            st.error("❌ Gemini API 키가 설정되지 않았습니다.")
        if st.button("🔄 모델 캐시 초기화", use_container_width=True):
            _model_cache.clear()
            st.toast("모델 캐시가 초기화되었습니다!")
            st.rerun()
        if st.button("🔄 DB 초기화"):
            if os.path.exists("academy.db"): os.remove("academy.db")
            init_database()
            st.rerun()

# ====================================================================
# 📊 메뉴 1: 대시보드 (통계 + 꺾은선 그래프)
# ====================================================================
if st.session_state.menu == "📊 대시보드":
    st.markdown("<div class='main-header'>📊 통합 대시보드</div>", unsafe_allow_html=True)
    
    stats = get_dashboard_stats()
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f"<div class='stat-card'><div class='stat-number'>{stats['total_students']}</div><div class='stat-label'>👥 총 학생</div></div>", unsafe_allow_html=True)
    with col2:
        st.markdown(f"<div class='stat-card'><div class='stat-number'>{stats['total_classes']}</div><div class='stat-label'>📚 총 반</div></div>", unsafe_allow_html=True)
    with col3:
        st.markdown(f"<div class='stat-card'><div class='stat-number'>{stats['total_schools']}</div><div class='stat-label'>🏫 총 학교</div></div>", unsafe_allow_html=True)
    with col4:
        st.markdown(f"<div class='stat-card'><div class='stat-number'>{stats['total_sessions']}</div><div class='stat-label'>🎯 총 잔여횟수</div></div>", unsafe_allow_html=True)
    
    st.divider()
    
    # 월별 학생 수 꺾은선 그래프
    st.markdown("<div class='sub-header'>📈 월별 학생 수 추이</div>", unsafe_allow_html=True)
    months = [f"2026-{i:02d}" for i in range(1, 13)]
    monthly_data = get_monthly_new_students()
    
    if monthly_data:
        df = pd.DataFrame(monthly_data)
        line = alt.Chart(df).mark_line(point=True, color="#D81B60", strokeWidth=3).encode(
            x=alt.X('month:T', title="월", axis=alt.Axis(format="%Y-%m")),
            y=alt.Y('cnt:Q', title="학생 수"),
            tooltip=['month', 'cnt']
        ).properties(height=300)
        st.altair_chart(line, use_container_width=True)
    else:
        st.info("📅 학생 데이터가 없습니다.")

# ====================================================================
# 📅 메뉴 2: 출석부 (달력형 + 반 폴더)
# ====================================================================
elif st.session_state.menu == "📅 출석부":
    st.markdown("<div class='main-header'>📅 출석부</div>", unsafe_allow_html=True)
    
    col_left, col_right = st.columns([1, 3])
    
    with col_left:
        st.markdown("#### 📅 월 선택")
        sel_year = st.selectbox("년도", ["2026", "2027", "2028"], index=0)
        sel_month = st.selectbox("월", [f"{i:02d}월" for i in range(1, 13)], index=datetime.now().month - 1)
        month_num = int(sel_month.replace("월", ""))
        year_month = f"{sel_year}-{sel_month:02d}" if isinstance(sel_month, int) else f"{sel_year}-{int(sel_month.replace('월','')):02d}"
        month_num = int(sel_month.replace("월", ""))
        year_month = f"{sel_year}-{month_num:02d}"
        
        st.divider()
        st.markdown("#### 📁 반 선택")
        all_c = get_all_classes()
        if all_c:
            for c in all_c:
                if st.button(f"📂 {c['name']} ({c['day_of_week']})", key=f"class_btn_{c['id']}"):
                    st.session_state.selected_class = c['id']
                    st.rerun()
        else:
            st.info("반을 먼저 등록해주세요.")
    
    with col_right:
        selected_class_id = st.session_state.get("selected_class")
        if selected_class_id and all_c:
            c_obj = next((c for c in all_c if c['id'] == selected_class_id), None)
            if c_obj:
                st.markdown(f"<div class='sub-header'>📂 {c_obj['name']} ({c_obj['day_of_week']} {c_obj['time_slot']})</div>", unsafe_allow_html=True)
                
                students = get_students_by_class(selected_class_id)
                if students:
                    # 개인 선택
                    student_names = {f"{s['name']} ({s.get('school_name','')})": s for s in students}
                    sel_student_name = st.selectbox("👤 학생 선택", list(student_names.keys()), key="sel_student_att")
                    s = student_names[sel_student_name]
                    
                    # 달력 표시
                    cal = calendar.monthcalendar(int(sel_year), month_num)
                    
                    st.markdown(f"**{s['name']}** | 📱학부모: {s['parent_phone']} | 🎯잔여 {s['remaining_sessions']}회")
                    
                    # 달력 그리드 (날짜 클릭 시 input_date 변경)
                    cols = st.columns(7)
                    day_names = ["월", "화", "수", "목", "금", "토", "일"]
                    for i, name in enumerate(day_names):
                        cols[i].markdown(f"<div style='text-align:center; font-weight:700; color:#AD1457;'>{name}</div>", unsafe_allow_html=True)
                    
                    # 미리 선택된 날짜 세션
                    if "cal_click_date" not in st.session_state:
                        st.session_state.cal_click_date = datetime.now().strftime("%Y-%m-%d")
                    
                    for week in cal:
                        cols = st.columns(7)
                        for i, day in enumerate(week):
                            if day == 0:
                                cols[i].markdown("<div class='calendar-day'></div>", unsafe_allow_html=True)
                            else:
                                date_str = f"{sel_year}-{month_num:02d}-{day:02d}"
                                att = get_attendance_by_date(s["id"], date_str)
                                status = att['status'] if att else "미입력"
                                
                                css_class = "calendar-day"
                                if status == "출석": css_class += " present"
                                elif status == "결석": css_class += " absent"
                                elif status == "지각": css_class += " late"
                                elif status == "영상보강": css_class += " video"
                                
                                emoji = {"출석": "✅", "결석": "❌", "지각": "⏰", "영상보강": "📺", "미입력": ""}
                                
                                # 날짜 클릭 버튼
                                if cols[i].button(f"{day}", key=f"cal_{s['id']}_{date_str}"):
                                    st.session_state.cal_click_date = date_str
                                    st.rerun()
                    
                    st.divider()
                    
                    # 날짜별 출석 입력 (달력 클릭으로 선택된 날짜 사용)
                    st.markdown("#### ✏️ 출석 입력/수정")
                    try:
                        default_date = datetime.strptime(st.session_state.cal_click_date, "%Y-%m-%d")
                    except:
                        default_date = datetime.now()
                    
                    input_date = st.date_input("날짜 선택 (달력 날짜 클릭 or 직접 선택)", value=default_date, key=f"att_input_date_{s['id']}")
                    input_date_str = input_date.strftime("%Y-%m-%d")
                    
                    cols = st.columns([3, 3, 2])
                    att = get_attendance_by_date(s["id"], input_date_str)
                    current_status = att['status'] if att else "미입력"
                    
                    cols[0].markdown(f"**{s['name']}**")
                    new_status = cols[1].selectbox(
                        "상태",
                        ["미입력", "출석", "결석", "지각", "영상보강"],
                        index=["미입력", "출석", "결석", "지각", "영상보강"].index(current_status),
                        key=f"att_{s['id']}_{input_date_str}"
                    )
                    if cols[2].button("저장", key=f"save_{s['id']}_{input_date_str}"):
                        mark_attendance(s['id'], new_status, input_date_str)
                        st.toast(f"{s['name']} {new_status} 저장완료!")
                        st.rerun()
                else:
                    st.info("해당 반에 등록된 학생이 없습니다.")
        else:
            st.info("👈 왼쪽에서 반을 선택해주세요.")

# ====================================================================
# ⚙️ 메뉴 3: 학원 관리
# ====================================================================
elif st.session_state.menu == "⚙️ 학원 관리":
    st.markdown("<div class='main-header'>⚙️ 학원 관리</div>", unsafe_allow_html=True)
    t1, t2, t3 = st.tabs(["👥 학생 관리", "📚 반 관리", "🏫 학교 관리"])
    
    all_c = get_all_classes()
    all_sch = get_all_schools()

    with t1:
        with st.expander("➕ 신규 학생 등록", expanded=True):
            with st.form("add_student_form"):
                c1, c2 = st.columns(2)
                name = c1.text_input("* 이름")
                phone = c1.text_input("학생 연락처")
                parent = c1.text_input("* 학부모 연락처")
                cls = c2.selectbox("* 반", [c["name"] for c in all_c])
                sch = c2.selectbox("학교", [""] + [s["name"] for s in all_sch])
                sess = c2.number_input("초기 잔여 횟수", value=8)
                if st.form_submit_button("등록하기"):
                    if name and parent and cls:
                        cid = next(c["id"] for c in all_c if c["name"] == cls)
                        sid = None
                        if sch:
                            sid = next((s["id"] for s in all_sch if s["name"] == sch), None)
                        add_student(cid, name, phone, sess, parent, sid)
                        st.toast(f"{name} 등록 완료!")
                        st.rerun()
                    else:
                        st.error("이름, 학부모 연락처, 반은 필수 입력입니다.")
        
        all_s = get_all_students()
        for s in all_s:
            with st.expander(f"✏️ {s['name']} | {s['school_name'] or '학교미지정'} | {s['class_name']} | 📱{s['parent_phone']} | 🎯{s['remaining_sessions']}회"):
                with st.form(f"edit_student_{s['id']}"):
                    col1, col2 = st.columns(2)
                    new_name = col1.text_input("이름", value=s['name'])
                    new_phone = col1.text_input("학생 연락처", value=s['phone'])
                    new_parent = col1.text_input("학부모 연락처", value=s['parent_phone'])
                    new_cls = col2.selectbox("반", [c["name"] for c in all_c], index=[c["name"] for c in all_c].index(s['class_name']) if s['class_name'] in [c["name"] for c in all_c] else 0)
                    new_sch = col2.selectbox("학교", [""] + [s["name"] for s in all_sch], index=([""] + [s["name"] for s in all_sch]).index(s['school_name']) if s['school_name'] in [""] + [s["name"] for s in all_sch] else 0)
                    new_sess = col2.number_input("잔여 횟수", value=s['remaining_sessions'], key=f"sess_{s['id']}")
                    
                    cb1, cb2 = st.columns(2)
                    if cb1.form_submit_button("💾 수정 저장"):
                        new_cid = next(c["id"] for c in all_c if c["name"] == new_cls)
                        new_sid = next((s["id"] for s in all_sch if s["name"] == new_sch), None) if new_sch else None
                        update_student(s['id'], new_name, new_phone, new_parent, new_cid, new_sid, new_sess)
                        st.toast(f"{new_name} 정보 수정 완료!")
                        st.rerun()
                    if cb2.form_submit_button("🗑️ 삭제"):
                        delete_student(s['id'])
                        st.rerun()

    with t2:
        with st.form("add_class_form"):
            name = st.text_input("반 이름")
            days = st.multiselect("요일", ["월", "화", "수", "목", "금", "토", "일"])
            h = st.selectbox("시간", [f"오전 {i}시" for i in range(1,13)] + [f"오후 {i}시" for i in range(1,13)])
            m = st.selectbox("분", [f"{i:02d}분" for i in range(0,60,5)])
            notes = st.text_area("비고")
            if st.form_submit_button("반 생성"):
                add_class(name, "/".join(days), f"{h} {m}", notes)
                st.rerun()
        for c in all_c:
            st.write(f"📚 {c['name']} ({c['day_of_week']} {c['time_slot']})")
            if st.button("삭제", key=f"del_c_{c['id']}"): 
                delete_class(c['id'])
                st.rerun()
        
        st.divider()
        st.markdown("#### 📝 반별 영역 관리 (시험 영역)")
        if all_c:
            sel_class_for_area = st.selectbox("영역 관리할 반 선택", [c['name'] for c in all_c], key="area_class_select")
            c_obj_area = next(c for c in all_c if c['name'] == sel_class_for_area)
            
            with st.form("add_area_form"):
                new_area_name = st.text_input("새 영역 이름", placeholder="예: 어휘, 문법, 독해, 듣기")
                if st.form_submit_button("➕ 영역 추가"):
                    if new_area_name.strip():
                        add_class_area(c_obj_area['id'], new_area_name.strip())
                        st.toast(f"영역 '{new_area_name.strip()}' 추가 완료!")
                        st.rerun()
            
            existing_areas = get_class_areas(c_obj_area['id'])
            if existing_areas:
                for area in existing_areas:
                    c1, c2 = st.columns([5, 1])
                    c1.write(f"  📌 {area['area_name']}")
                    if c2.button("🗑️", key=f"del_area_{area['id']}"):
                        delete_class_area(area['id'])
                        st.toast(f"영역 '{area['area_name']}' 삭제 완료!")
                        st.rerun()
            else:
                st.info("영역이 없습니다. 위에서 추가해주세요.")
        else:
            st.info("반을 먼저 등록해주세요.")

    with t3:
        new_sch = st.text_input("학교명")
        if st.button("학교 추가"): 
            add_school(new_sch)
            st.rerun()
        for s in all_sch:
            st.write(f"🏫 {s['name']}")
            if st.button("삭제", key=f"del_sch_{s['id']}"): 
                delete_school(s['id'])
                st.rerun()

# ====================================================================
# 📖 메뉴 4: 교재분석
# ====================================================================
elif st.session_state.menu == "📖 교재분석":
    st.markdown("<div class='main-header'>📖 교재 분석 & 단어장 생성</div>", unsafe_allow_html=True)
    file = st.file_uploader("교재 업로드 (PDF/이미지)", type=["pdf", "png", "jpg", "jpeg"])
    if file:
        if st.button("🤖 AI 100개 핵심 단어 추출", type="primary", use_container_width=True):
            vocab_log = st.empty()
            with st.spinner("AI 교재 분석 중..."):
                api_key = st.session_state.gemini_key_input
                if not api_key:
                    st.error("설정에서 Gemini API Key를 입력해주세요.")
                else:
                    try:
                        vocab_log.info("⏳ 1/3 📄 교재 파일 텍스트 추출 중...")
                        if file.name.lower().endswith('.pdf'):
                            text, success = process_pdf_with_ocr(api_key, file.getvalue())
                        else:
                            text, success = extract_text_from_image(api_key, file.getvalue())
                        
                        if success:
                            vocab_log.info("⏳ 2/3 🤖 AI가 핵심 단어 추출 중...")
                            result = extract_vocabulary(api_key, text)
                            
                            if isinstance(result, dict) and "error" in result:
                                vocab_log.error(f"❌ 단어 추출 실패: {result['error']}")
                                st.error(f"⚠️ {result['error']}")
                            elif result:
                                st.session_state.vocab_result = result
                                vocab_log.info(f"⏳ 3/3 📊 {len(result)}개 단어 정리 중...")
                                vocab_log.success(f"✅ {len(result)}개 단어 추출 완료!")
                                st.toast("분석 완료!")
                            else:
                                vocab_log.warning("⚠️ 단어를 추출하지 못했습니다. 텍스트 내용을 확인해주세요.")
                                st.warning("단어를 추출하지 못했습니다.")
                        else:
                            vocab_log.error(f"❌ 텍스트 추출 실패: {text}")
                            st.error(f"텍스트 추출 실패: {text}")
                    except Exception as e:
                        vocab_log.error(f"❌ 분석 중 오류: {e}")
                        st.error(f"분석 중 오류: {e}")
                st.rerun()

    if st.session_state.vocab_result:
        df = pd.DataFrame(st.session_state.vocab_result)
        st.markdown("<div class='sub-header'>📋 단어장 미리보기</div>", unsafe_allow_html=True)
        selected = []
        for i, row in df.iterrows():
            c1, c2, c3 = st.columns([0.5, 2, 8])
            if c1.checkbox(f"{i+1}", key=f"v_{i}", value=True): selected.append(i)
            c2.markdown(f"**{row['word']}**")
            derivs = row.get('derivatives', '')
            synos = row.get('synonyms', '')
            info_parts = [row['meaning']]
            if isinstance(derivs, list) and derivs:
                info_parts.append(f"<span style='color:#888; font-size:0.85rem;'>파생어: {', '.join(str(d) for d in derivs if d)}</span>")
            elif isinstance(derivs, str) and derivs:
                info_parts.append(f"<span style='color:#888; font-size:0.85rem;'>파생어: {derivs}</span>")
            if isinstance(synos, list) and synos:
                info_parts.append(f"<span style='color:#AD1457; font-size:0.85rem;'>유의어: {', '.join(str(s) for s in synos if s)}</span>")
            elif isinstance(synos, str) and synos:
                info_parts.append(f"<span style='color:#AD1457; font-size:0.85rem;'>유의어: {synos}</span>")
            c3.markdown("<br>".join(info_parts), unsafe_allow_html=True)
        
        if st.button("📥 파일 생성 및 다운로드"):
            with st.spinner("파일 생성 중..."):
                # 파생어/유의어를 문자열로 변환하는 헬퍼
                def to_str(val):
                    if isinstance(val, list):
                        return ", ".join(str(v) for v in val if v)
                    return str(val) if val else ""
                
                # 엑셀용 데이터프레임 (파생어, 유의어 별도 컬럼)
                excel_data = []
                for _, r in df.iloc[selected].iterrows():
                    excel_data.append({
                        "단어": r['word'],
                        "뜻": r['meaning'],
                        "파생어": to_str(r.get('derivatives', '')),
                        "유의어": to_str(r.get('synonyms', ''))
                    })
                excel_df = pd.DataFrame(excel_data)
                output_ex = io.BytesIO()
                with pd.ExcelWriter(output_ex, engine='openpyxl') as writer: 
                    excel_df.to_excel(writer, index=False)
                
                # 워드용 (파생어, 유의어 별도 컬럼)
                doc = Document()
                doc.add_heading('🎀 대치앨리영어 단어장', 0)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                for h_idx, h_name in enumerate(['단어', '뜻', '파생어', '유의어']):
                    table.rows[0].cells[h_idx].text = h_name
                for _, r in df.iloc[selected].iterrows():
                    cells = table.add_row().cells
                    cells[0].text = str(r['word'])
                    cells[1].text = str(r['meaning'])
                    cells[2].text = to_str(r.get('derivatives', ''))
                    cells[3].text = to_str(r.get('synonyms', ''))
                output_wd = io.BytesIO(); doc.save(output_wd)
            st.success("생성 완료!")
            st.download_button("📥 Excel 다운로드", output_ex.getvalue(), "앨리단어장.xlsx")
            st.download_button("📥 Word 다운로드", output_wd.getvalue(), "앨리단어장.docx")

# ====================================================================
# 🔬 메뉴 5: 시험지 OCR
# ====================================================================
elif st.session_state.menu == "🔬 시험지 OCR":
    st.markdown("<div class='main-header'>🔬 시험지 OCR 분석</div>", unsafe_allow_html=True)
    all_s = get_all_students()
    if all_s:
        c1, c2 = st.columns(2)
        student_map = {f"{s['name']} ({s['school_name'] or '미지정'})": s for s in all_s}
        sel_s = c1.selectbox("학생 선택", list(student_map.keys()))
        week = c2.selectbox("주차", [f"{i}주차" for i in range(1, 21)])
        
        files = st.file_uploader("시험지 업로드 (다중 가능)", accept_multiple_files=True, type=["pdf", "png", "jpg", "jpeg"])
        if files:
            with st.expander("🖼️ 이미지 확인"):
                for f in files: 
                    if not f.name.endswith(".pdf"): st.image(f, use_column_width=True)
            
            if st.button("🔍 AI 자동 채점", type="primary"):
                # 진행 로그 표시를 위한 placeholder
                log_placeholder = st.empty()
                with st.spinner("AI 분석 중..."):
                    api_key = st.session_state.gemini_key_input
                    if not api_key:
                        st.error("설정에서 Gemini API Key를 입력해주세요.")
                    else:
                        # 진행 로그
                        log_placeholder.info("⏳ 1/3 📤 파일 업로드 및 텍스트 변환 중...")
                        f_info = [{"type": "pdf" if f.name.endswith(".pdf") else "image", "bytes": f.getvalue()} for f in files]
                        
                        log_placeholder.info("⏳ 2/3 🤖 Gemini AI 분석 요청 중...")
                        result = analyze_exam_multi(api_key, f_info, week)
                        
                        if result and "error" not in result:
                            log_placeholder.info("⏳ 3/3 📊 결과 데이터 처리 중...")
                            st.session_state.last_exam_result = result
                            st.session_state.ocr_feedback = [True] * result.get('total_questions', 0)
                            log_placeholder.success("✅ AI 분석 완료!")
                            st.toast("분석 성공!")
                        else:
                            err_msg = result.get("error") if result else "알 수 없는 오류"
                            log_placeholder.error(f"⚠️ 분석 실패: {err_msg}")
                            st.error(f"⚠️ 분석 실패: {err_msg}")

        if st.session_state.last_exam_result:
            res = st.session_state.last_exam_result
            if "error" in res:
                st.error(res["error"])
                st.session_state.last_exam_result = None
            else:
                total_q = res.get('total_questions', 0)
                area_acc = res.get('area_accuracy_percent', {})
                area_qs = res.get('area_questions', {})
                weakness = res.get('weakness_analysis', '')
                prescription = res.get('prescription', '')
                
                # ── 상단 요약 카드 ──
                st.markdown("<div class='sub-header'>📊 AI 채점 결과</div>", unsafe_allow_html=True)
                
                summary_cols = st.columns(5)
                pct = int(res.get('correct_answers', 0) / total_q * 100) if total_q > 0 else 0
                color = "#4CAF50" if pct >= 80 else "#FF9800" if pct >= 60 else "#F44336"
                
                with summary_cols[0]:
                    st.markdown(f"""
                    <div style="background:white; padding:1.2rem; border-radius:15px; border:2px solid {color}; text-align:center;">
                        <div style="font-size:2.5rem; font-weight:800; color:{color};">{pct}%</div>
                        <div style="font-size:0.8rem; color:#888;">총 정답률</div>
                        <div style="font-size:0.75rem; color:#666;">{res.get('correct_answers',0)}/{total_q}문항</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                area_icons = {"어휘": "📝", "문법": "📐", "독해": "📖", "듣기": "🎧"}
                for i, (area, acc) in enumerate(area_acc.items()):
                    if i >= 4: break
                    q_cnt = area_qs.get(area, '?')
                    icon = area_icons.get(area, "📋")
                    a_color = "#4CAF50" if acc >= 80 else "#FF9800" if acc >= 60 else "#F44336"
                    with summary_cols[i + 1]:
                        st.markdown(f"""
                        <div style="background:white; padding:1.2rem; border-radius:15px; border:2px solid {a_color}20; text-align:center; border-left: 4px solid {a_color};">
                            <div style="font-size:0.9rem;">{icon} {area}</div>
                            <div style="font-size:1.8rem; font-weight:800; color:{a_color};">{acc}%</div>
                            <div style="font-size:0.75rem; color:#888;">{q_cnt}문항</div>
                        </div>
                        """, unsafe_allow_html=True)
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                # ── 영역별 정답률 게이지 ──
                st.markdown("#### 📈 영역별 정답률")
                for area, acc in area_acc.items():
                    icon = area_icons.get(area, "📋")
                    a_color = "#4CAF50" if acc >= 80 else "#FF9800" if acc >= 60 else "#F44336"
                    st.markdown(f"""
                    <div style="display:flex; align-items:center; gap:10px; margin-bottom:8px;">
                        <span style="min-width:70px; font-weight:600;">{icon} {area}</span>
                        <div style="flex:1; background:#f0f0f0; border-radius:10px; height:24px;">
                            <div style="background:{a_color}; width:{acc}%; height:24px; border-radius:10px; display:flex; align-items:center; justify-content:flex-end; padding-right:8px;">
                                <span style="color:white; font-size:0.8rem; font-weight:700;">{acc}%</span>
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                # ── 문제별 정답 토글 ──
                st.markdown("#### ✅ 문제별 정답 확인")
                st.caption("각 문제를 클릭하여 정답(ON)/오답(OFF)을 토글하세요.")
                
                questions = res.get('questions', [])
                new_feedback = []
                
                if questions:
                    # AI가 문제별 상세 정보를 제공한 경우
                    area_colors = {"어휘": "#E91E63", "문법": "#9C27B0", "독해": "#2196F3", "듣기": "#FF9800"}
                    
                    # 영역별로 그룹화
                    area_groups = {}
                    for q in questions:
                        a = q.get('area', '기타')
                        if a not in area_groups:
                            area_groups[a] = []
                        area_groups[a].append(q)
                    
                    for area, area_questions_list in area_groups.items():
                        acc = area_acc.get(area, 0)
                        a_color = area_colors.get(area, "#888")
                        st.markdown(f"""
                        <div style="background:{a_color}10; border-left:4px solid {a_color}; padding:8px 12px; border-radius:0 8px 8px 0; margin:12px 0 6px;">
                            <span style="font-weight:700; color:{a_color};">{area_icons.get(area,'')} {area}</span>
                            <span style="font-size:0.85rem; color:#666;"> ({len(area_questions_list)}문항 | 정답률 {acc}%)</span>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        for q in area_questions_list:
                            q_num = q.get('number', 0)
                            is_correct_val = q.get('correct', True)
                            q_type = q.get('type', '')
                            q_summary = q.get('summary', '')
                            
                            c1, c2, c3 = st.columns([0.5, 3, 7])
                            with c1:
                                is_correct = st.toggle(
                                    f"Q{q_num}",
                                    value=is_correct_val,
                                    key=f"ox_{q_num}",
                                    help=f"{area} 영역 - 정답 토글"
                                )
                                new_feedback.append(is_correct)
                            with c2:
                                st.markdown(f"**{q_type}**")
                            with c3:
                                st.markdown(f"<span style='color:#666; font-size:0.85rem;'>{q_summary}</span>", unsafe_allow_html=True)
                else:
                    # 문제별 상세 정보 없으면 기본 표시
                    cols = st.columns(5)
                    for i in range(total_q):
                        with cols[i % 5]:
                            is_correct = st.toggle(
                                f"Q{i + 1}",
                                value=True,
                                key=f"ox_{i}",
                                help="ON=정답, OFF=오답"
                            )
                            new_feedback.append(is_correct)
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                # ── 최종 정답률 카드 ──
                correct = sum(new_feedback)
                final_pct = int(correct / total_q * 100) if total_q > 0 else 0
                final_color = "#4CAF50" if final_pct >= 80 else "#FF9800" if final_pct >= 60 else "#F44336"
                
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, {final_color}10, {final_color}05); padding:1.5rem; border-radius:15px; border:2px solid {final_color}40; margin:1rem 0;">
                    <div style="display:flex; justify-content:space-between; align-items:center;">
                        <div>
                            <span style="font-size:1.3rem; font-weight:700;">📊 최종 정답률</span>
                            <span style="font-size:0.85rem; color:#888;"> (수동 수정 반영)</span>
                        </div>
                        <div style="text-align:right;">
                            <span style="font-size:2.2rem; font-weight:800; color:{final_color};">{final_pct}%</span>
                            <span style="font-size:1rem; color:#888;"> ({correct}/{total_q})</span>
                        </div>
                    </div>
                    <div style="background:#f0f0f0; border-radius:10px; height:18px; margin-top:10px;">
                        <div style="background:{final_color}; width:{final_pct}%; height:18px; border-radius:10px; transition:width 0.5s;"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # ── AI 분석 요약 ──
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"""
                    <div style="background:#FFF5F8; padding:1rem; border-radius:12px; border-left:4px solid #E91E63; margin-bottom:1rem;">
                        <div style="font-weight:700; color:#AD1457; margin-bottom:5px;">🔍 취약점 분석</div>
                        <div style="font-size:0.9rem; color:#444;">{weakness}</div>
                    </div>
                    """, unsafe_allow_html=True)
                with col2:
                    st.markdown(f"""
                    <div style="background:#F3E5F5; padding:1rem; border-radius:12px; border-left:4px solid #9C27B0; margin-bottom:1rem;">
                        <div style="font-weight:700; color:#7B1FA2; margin-bottom:5px;">💊 학습 처방</div>
                        <div style="font-size:0.9rem; color:#444;">{prescription}</div>
                    </div>
                    """, unsafe_allow_html=True)
                
                # ── 액션 버튼 ──
                col1, col2 = st.columns(2)
                if col1.button("💾 성적으로 저장", type="primary", use_container_width=True):
                    s_obj = student_map[sel_s]
                    add_student_history(s_obj["id"], "성적", f"[{week}] {correct}/{total_q} 정답({final_pct}%). {weakness}", week)
                    add_score(s_obj["id"], 
                              area_acc.get('듣기', 0),
                              area_acc.get('어휘', 0),
                              area_acc.get('문법', 0),
                              area_acc.get('독해', 0),
                              f"{week} OCR")
                    st.toast("성적에 저장 완료!")
                    st.rerun()
                
                if col2.button("📤 학부모님께 문자 전송", use_container_width=True):
                    s_obj = student_map[sel_s]
                    msg = f"[앨리영어] {s_obj['name']} {week} 시험결과: {correct}/{total_q}({final_pct}%)\n취약: {weakness}\n처방: {prescription}"
                    send_message(st.session_state.solapi_key_input, st.session_state.solapi_secret_input, s_obj['parent_phone'], st.session_state.solapi_from_input, msg)
                    st.toast("문자 전송 완료!")
    else:
        st.info("학생을 먼저 등록해주세요.")

# ====================================================================
# 📈 메뉴 6: 성적 & 리포트
# ====================================================================
elif st.session_state.menu == "📈 성적 & 리포트":
    st.markdown("<div class='main-header'>📈 성적 관리 및 월간 리포트</div>", unsafe_allow_html=True)
    all_s = get_all_students()
    if all_s:
        sel_s_name = st.selectbox("학생 선택", [f"{s['name']} ({s['school_name'] or '미지정'})" for s in all_s])
        s_obj = next(s for s in all_s if f"{s['name']} ({s['school_name'] or '미지정'})" == sel_s_name)
        
        class_areas_list = get_class_areas(s_obj.get('class_id', 0)) if s_obj.get('class_id') else []
        area_names = [a['area_name'] for a in class_areas_list] if class_areas_list else ["듣기", "어휘", "문법", "독해"]
        
        st.markdown("#### 📝 새 성적 입력")
        with st.form("score_form"):
            c1, c2 = st.columns(2)
            tn = c1.text_input("시험명", value=f"{datetime.now().strftime('%m월 %d일')} 테스트")
            dt = c1.date_input("날짜")
            area_scores = {}
            for i, area_name in enumerate(area_names):
                with (c1 if i < len(area_names) // 2 else c2):
                    area_scores[area_name] = st.text_input(area_name, placeholder="예: 5/9", key=f"new_{area_name}")
            if st.form_submit_button("저장"):
                score_details = ", ".join([f"{k}: {v}" for k, v in area_scores.items() if v])
                add_score(s_obj['id'], 0, 0, 0, 0, tn, dt.strftime("%Y-%m-%d"), score_details)
                st.toast("성적 저장 완료!")
                st.rerun()
        st.divider()
        st.markdown("#### ✏️ 기존 성적 수정/삭제")
        scores = get_student_scores(s_obj['id'])
        if scores:
            for sc in scores:
                display_text = sc.get('score_details', '') or f"합계:{sc['listening']}"
                with st.expander(f"📅 {sc['date']} - {sc['test_name']} ({display_text})"):
                    with st.form(f"edit_score_{sc['id']}"):
                        c1, c2 = st.columns(2)
                        new_tn = c1.text_input("시험명", value=sc['test_name'])
                        new_dt = c1.date_input("날짜", value=datetime.strptime(sc['date'], "%Y-%m-%d") if sc['date'] else datetime.now())
                        old_details = sc.get('score_details', '')
                        old_scores = {}
                        if old_details:
                            for part in old_details.split(','):
                                if ':' in part:
                                    k, v = part.split(':', 1)
                                    old_scores[k.strip()] = v.strip()
                        new_scores = {}
                        for i, area_name in enumerate(area_names):
                            default_val = old_scores.get(area_name, '')
                            with (c1 if i < len(area_names) // 2 else c2):
                                new_scores[area_name] = st.text_input(area_name, value=default_val, key=f"edit_{sc['id']}_{area_name}")
                        new_details = ", ".join([f"{k}: {v}" for k, v in new_scores.items() if v])
                        cb1, cb2 = st.columns(2)
                        if cb1.form_submit_button("💾 수정 저장"):
                            update_score(sc['id'], 0, 0, 0, 0, new_tn, new_details)
                            st.rerun()
                        if cb2.form_submit_button("🗑️ 삭제"):
                            delete_score(sc['id'])
                            st.rerun()
        else:
            st.info("성적 기록이 없습니다.")
    else:
        st.info("학생이 없습니다.")

# ====================================================================
# 💰 메뉴 7: 횟수 관리
# ====================================================================
elif st.session_state.menu == "💰 횟수 관리":
    st.markdown("<div class='main-header'>💰 횟수 관리</div>", unsafe_allow_html=True)
    
    all_s = get_all_students()
    if all_s:
        student_map = {f"{s['name']} ({s['school_name'] or '미지정'})": s for s in all_s}
        sel_name = st.selectbox("학생 선택", list(student_map.keys()))
        s = student_map[sel_name]
        
        # 예상 만료일 계산 (주 2회 기준: 4회당 약 2주 소모)
        remaining = s['remaining_sessions']
        if remaining > 0:
            weeks_left = max(1, round(remaining / 2))
            estimated_end = (datetime.now() + timedelta(weeks=weeks_left)).strftime("%Y-%m-%d")
        else:
            estimated_end = "⚠️ 소진됨"
        
        # 학생 정보
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("현재 잔여횟수", f"{s['remaining_sessions']}회")
        col2.metric("예상 만료일 (주2회 기준)", estimated_end)
        
        with col3:
            # 충전 입력
            charge_amount = st.number_input("충전 횟수", min_value=1, max_value=100, value=8)
            charge_note = st.text_input("충전 메모 (예: 5월 수강권)")
            if st.button("💰 충전", type="primary"):
                new_balance = add_payment_transaction(s['id'], "charge", charge_amount, f"[충전] {charge_note or f'{charge_amount}회'}")
                st.toast(f"{charge_amount}회 충전 완료! (잔여 {new_balance}회)")
                st.rerun()
        
        with col4:
            # 다음 결제일 설정
            next_pay = st.date_input("다음 결제일 설정")
            if st.button("📅 저장"):
                set_next_payment_date(s['id'], next_pay.strftime("%Y-%m-%d"))
                st.toast("결제일 저장 완료!")
                st.rerun()
        
        st.divider()
        
        st.markdown("#### 📋 거래 내역")
        transactions = get_payment_transactions(s['id'])
        if transactions:
            for t in transactions:
                c1, c2, c3, c4, c5, c6 = st.columns([2, 1, 1.5, 1.5, 3, 1])
                c1.write(t['date'])
                c2.write("💰" if t['type'] == 'charge' else "📉")
                c3.write(f"+{t['amount']}회" if t['amount'] > 0 else f"{t['amount']}회")
                c4.write(f"{t['balance_after']}회")
                c5.write(t['note'])
                if c6.button("🗑️", key=f"del_tx_{t['id']}"):
                    delete_payment_transaction(t['id'])
                    st.toast("거래 내역 삭제 완료!")
                    st.rerun()
        else:
            st.info("거래 내역이 없습니다.")
        
        # 예상 차감 내역 (이번달 출석 기준)
        st.divider()
        st.markdown("#### 📅 예상 차감 내역 (이번달)")
        this_month = datetime.now().strftime("%Y-%m")
        atts = get_month_attendance(s['id'], this_month)
        if atts:
            deducted = sum(1 for a in atts if a['status'] != "미입력")
            st.info(f"이번달 {len(atts)}회 중 {deducted}회 차감 완료 (잔여 {s['remaining_sessions'] - deducted}회 예상)")
        else:
            st.info("이번달 출석 기록이 없습니다.")
    else:
        st.info("학생이 없습니다.")

# ====================================================================
# 🏫 메뉴 8: 학교별 기출
# ====================================================================
elif st.session_state.menu == "🏫 학교별 기출":
    st.markdown("<div class='main-header'>🏫 학교별 기출 분석</div>", unsafe_allow_html=True)
    t1, t2 = st.tabs(["🔍 대시보드", "📤 업로드 및 분석"])
    with t2:
        with st.form("exam_up"):
            sch = st.selectbox("학교", [s['name'] for s in get_all_schools()])
            g = st.selectbox("학년", ["중1","중2","중3","고1","고2","고3"])
            sem_options = ["1학기 중간", "1학기 기말", "2학기 중간", "2학기 기말", "기타 (직접입력)"]
            sem = st.selectbox("학기", sem_options)
            if sem == "기타 (직접입력)":
                sem = st.text_input("학기 직접 입력", value="")
            files = st.file_uploader("기출 파일 업로드 (다중 파일 가능 - PDF/이미지)", accept_multiple_files=True)
            if st.form_submit_button("🤖 AI 분석 및 저장"):
                if files and st.session_state.gemini_key_input:
                    with st.spinner(f"AI가 {len(files)}개의 파일을 분석 중입니다..."):
                        # 파일 정보 수집
                        all_text = ""
                        for f in files:
                            if f.name.lower().endswith('.pdf'):
                                text, _ = process_pdf_with_ocr(st.session_state.gemini_key_input, f.getvalue())
                            else:
                                text, _ = extract_text_from_image(st.session_state.gemini_key_input, f.getvalue())
                            all_text += text + "\n\n"
                        
                        # AI 분석
                        prompt = f"다음은 {sch} {g} {sem} 기출문제 분석입니다. 문제 출제 경향, 난이도, 주요 포인트를 3~5줄로 분석해주세요:\n\n{all_text[:5000]}"
                        try:
                            response, model_name = _try_models_generate(
                                get_client(st.session_state.gemini_key_input), 
                                prompt,
                                st.session_state.gemini_key_input
                            )
                            ai_report = response.text
                        except Exception as e:
                            ai_report = f"AI 분석 실패: {str(e)}"
                        
                        save_school_exam(sch, g, "2024", sem, "", ai_report)
                        st.success("분석 및 저장 완료!")
                        st.rerun()
                else:
                    st.warning("파일을 업로드하고 Gemini API Key를 설정해주세요.")
    
    with t1:
        schs = get_distinct_schools_from_exams()
        if schs:
            sel = st.selectbox("학교 선택", schs)
            exams = get_school_exams(sel)
            for e in exams:
                with st.expander(f"📅 {e['year']} {e['semester']} ({e['grade']})"):
                    st.markdown(f"""
                    <div style="background: #FFF5F8; padding: 1rem; border-radius: 10px; border-left: 4px solid #D81B60;">
                        <b>🤖 AI 분석 리포트</b><br>
                        {e['ai_report']}
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("등록된 분석 자료가 없습니다. [📤 업로드 및 분석] 탭에서 기출문제를 업로드해주세요.")

# ====================================================================
# 📋 메뉴 9: 히스토리
# ====================================================================
elif st.session_state.menu == "📋 히스토리":
    st.markdown("<div class='main-header'>📋 전체 히스토리</div>", unsafe_allow_html=True)
    all_s = get_all_students()
    if all_s:
        sel_s = st.selectbox("학생 선택", [f"{s['name']} ({s['school_name'] or '미지정'})" for s in all_s])
        s_obj = next(s for s in all_s if f"{s['name']} ({s['school_name'] or '미지정'})" == sel_s)
        
        histories = get_student_history(s_obj['id'])
        if histories:
            for h in histories:
                with st.container():
                    st.markdown(f"""
                    <div class='info-card'>
                        <span style='color:#D81B60; font-weight:700;'>{h['date']}</span>
                        <span style='background:#FCE4EC; padding:2px 8px; border-radius:10px; font-size:0.8rem;'>{h['category']}</span>
                        {f"<span style='color:#888;'>({h['week_label']})</span>" if h['week_label'] else ""}
                        <p style='margin-top:0.5rem;'>{h['notes']}</p>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("히스토리 기록이 없습니다. 시험지 OCR 분석 결과나 성적을 저장하면 여기에 기록됩니다.")
    else:
        st.info("학생이 없습니다.")

# ====================================================================
# 📨 메뉴 10: 알림톡 발송
# ====================================================================
elif st.session_state.menu == "📨 알림톡 발송":
    st.markdown("<div class='main-header'>📨 카카오 알림톡 발송</div>", unsafe_allow_html=True)
    
    t1, t2, t3 = st.tabs(["📤 개별 발송", "📋 일괄 발송 (반별)", "📝 미리보기"])
    
    all_students = get_all_students()
    api_key = st.session_state.solapi_key_input
    api_secret = st.session_state.solapi_secret_input
    from_number = st.session_state.solapi_from_input
    
    with t1:
        st.markdown("### 📋 주간 Report 발송")
        if all_students:
            student_map = {f"{s['name']} ({s['school_name'] or '미지정'})": s for s in all_students}
            sel_name = st.selectbox("학생 선택", list(student_map.keys()), key="alim_student")
            s = student_map[sel_name]
            
            parent_phone = st.text_input("📱 학부모 연락처", value=s.get('parent_phone', ''), placeholder="01012345678")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**학생명:** {s['name']}")
                st.markdown(f"**학교:** {s['school_name'] or '미지정'}")
                st.markdown(f"**반:** {s.get('class_name', '')}")
            with col2:
                st.markdown(f"**학부모 연락처:** {s.get('parent_phone', '미등록')}")
            
            st.divider()
            
            st.markdown("#### 📅 이번 주 출석 현황")
            week_att = get_week_attendance(s['id'])
            if week_att:
                for att in week_att:
                    emoji = {"출석": "✅", "결석": "❌", "지각": "⏰", "영상보강": "📺", "미입력": "❓"}
                    st.markdown(f"**{att['date']} ({att['day_name']})**: {emoji.get(att['status'], '❓')} {att['status']}")
            else:
                st.info("이번 주 수업일이 없습니다.")
            
            st.divider()
            
            st.markdown("#### 📝 시험결과 (영역별 점수)")
            c_obj = next((c for c in get_all_classes() if c['name'] == s.get('class_name', '')), None)
            if c_obj:
                class_areas = get_class_areas(c_obj['id'])
                with st.form("add_area_form_t1", clear_on_submit=True):
                    new_area = st.text_input("새 영역 이름", placeholder="예: 어휘, 문법, 독해, 듣기")
                    if st.form_submit_button("➕ 영역 추가"):
                        if new_area.strip():
                            add_class_area(c_obj['id'], new_area.strip())
                            st.toast(f"영역 '{new_area.strip()}' 추가 완료!")
                            st.rerun()
                area_scores = {}
                if class_areas:
                    cols = st.columns(min(len(class_areas), 4))
                    for i, area in enumerate(class_areas):
                        with cols[i % len(cols)]:
                            score = st.text_input(area['area_name'], placeholder="예: 5/9", key=f"area_score_{area['id']}")
                            area_scores[area['area_name']] = score
                    del_cols = st.columns(min(len(class_areas), 4))
                    for i, area in enumerate(class_areas):
                        with del_cols[i % len(del_cols)]:
                            if st.button(f"🗑️ {area['area_name']}", key=f"del_area_{area['id']}"):
                                delete_class_area(area['id'])
                                st.toast(f"영역 '{area['area_name']}' 삭제 완료!")
                                st.rerun()
                else:
                    st.info("영역을 추가해주세요.")
            else:
                st.info("반 정보를 찾을 수 없습니다.")
            
            st.divider()
            
            st.markdown("#### ⚙️ 카카오 알림톡 설정")
            c1, c2 = st.columns(2)
            pf_id = c1.text_input("카카오 채널 ID (pfId)", value="KA01PF260529142530005vhA1Xuvwf5K", key="t1_pf")
            template_id = c2.text_input("템플릿 ID", value="KA01TP260529143311427H1YO1TIAFCF", key="t1_tpl")
            
            st.markdown("#### 📱 발송 내용 미리보기")
            preview_parts = [f"안녕하세요, {s['name']}학부모님.", "자신감을 키우는 대치앨리영어교육입니다.", ""]
            preview_parts.append("📅 이번 주 출석 현황")
            if week_att:
                for att in week_att:
                    preview_parts.append(f"• {att['date']} ({att['day_name']}): {att['status']}")
            preview_parts.append("")
            preview_parts.append("📝 시험결과")
            if c_obj:
                for area in get_class_areas(c_obj['id']):
                    sv = st.session_state.get(f"area_score_{area['id']}", "")
                    if sv:
                        preview_parts.append(f"• {area['area_name']}: {sv}")
            preview_parts.append("")
            preview_parts.append("❤️ 대치앨리영어")
            st.code("\n".join(preview_parts))
            
            if st.button("📨 발송하기", type="primary", use_container_width=True):
                if not parent_phone:
                    st.error("⚠️ 학부모 연락처를 입력해주세요.")
                elif not api_key or api_key in ("", "YOUR_SOLAPI_KEY"):
                    st.error("⚠️ 설정에서 Solapi API Key를 입력해주세요.")
                else:
                    msg_parts = [f"안녕하세요, {s['name']}학부모님.", "자신감을 키우는 대치앨리영어교육입니다.", ""]
                    msg_parts.append("📅 이번 주 출석 현황")
                    if week_att:
                        for att in week_att:
                            msg_parts.append(f"• {att['date']} ({att['day_name']}): {att['status']}")
                    msg_parts.append("")
                    msg_parts.append("📝 시험결과")
                    if c_obj:
                        for area in get_class_areas(c_obj['id']):
                            sv = st.session_state.get(f"area_score_{area['id']}", "")
                            if sv:
                                msg_parts.append(f"• {area['area_name']}: {sv}")
                    msg_parts.append("")
                    msg_parts.append("❤️ 대치앨리영어")
                    msg = "\n".join(msg_parts)
                    
                    if template_id:
                        att_str = " | ".join([f"{a['date']} ({a['day_name']}): {a['status']}" for a in week_att]) if week_att else ""
                        test_str = ""
                        if c_obj:
                            test_lines = []
                            for area in get_class_areas(c_obj['id']):
                                sv = st.session_state.get(f"area_score_{area['id']}", "")
                                if sv:
                                    test_lines.append(f"{area['area_name']}: {sv}")
                            test_str = " | ".join(test_lines)
                        variables = {
                            "#{학생이름}": s['name'], "#{출석날짜}": att_str, "#{출석상태}": "",
                            "#{잔여횟수}": "", "#{시험주차}": "", "#{정답률}": test_str,
                            "#{취약영역}": "", "#{리포트월}": "", "#{듣기점수}": "", "#{어휘점수}": "",
                            "#{문법점수}": "", "#{독해점수}": "", "#{AI총평}": "",
                        }
                        res = send_comprehensive_alimtalk(api_key, api_secret, from_number, parent_phone, pf_id, template_id, variables)
                    else:
                        res = send_message(api_key, api_secret, parent_phone, from_number, msg)
                    st.session_state.last_send_result = res
                    st.rerun()
            
            # 발송 결과
            if "last_send_result" in st.session_state:
                res = st.session_state.last_send_result
                if res.get("success"):
                    st.success(res.get("message", "✅ 발송 성공!"))
                elif res.get("is_test"):
                    st.info("🔧 테스트 모드 - API 키를 설정해주세요")
                    with st.expander("📝 테스트 메시지 내용"):
                        st.text(res.get("message", ""))
                else:
                    st.error(res.get("message", "❌ 발송 실패"))
    
    with t2:
        st.markdown("### 📚 반별 일괄 주간 Report 발송")
        all_c = get_all_classes()
        if all_c:
            sel_class = st.selectbox("반 선택", [c['name'] for c in all_c], key="batch_class")
            c_obj = next(c for c in all_c if c['name'] == sel_class)
            students_in_class = get_students_by_class(c_obj['id'])
            st.markdown(f"**학생 수:** {len(students_in_class)}명")
            class_areas_batch = get_class_areas(c_obj['id'])
            st.markdown(f"**등록된 영역:** {', '.join([a['area_name'] for a in class_areas_batch]) if class_areas_batch else '없음'}")
            
            c1, c2 = st.columns(2)
            batch_pf_id = c1.text_input("카카오 채널 ID", value="KA01PF260529142530005vhA1Xuvwf5K", key="batch_pf")
            batch_template_id = c2.text_input("템플릿 ID", value="KA01TP260529143311427H1YO1TIAFCF", key="batch_tpl")
            
            st.markdown("##### 📱 발송 내용 미리보기 (첫번째 학생 기준)")
            if students_in_class:
                preview_s = students_in_class[0]
                week_att_p = get_week_attendance(preview_s['id'])
                pp = [f"안녕하세요, {preview_s['name']}학부모님.", "자신감을 키우는 대치앨리영어교육입니다.", "", "📅 이번 주 출석 현황"]
                if week_att_p:
                    for a in week_att_p:
                        pp.append(f"• {a['date']} ({a['day_name']}): {a['status']}")
                pp += ["", "📝 시험결과"]
                if class_areas_batch:
                    for a in class_areas_batch:
                        pp.append(f"• {a['area_name']}: -")
                pp += ["", "❤️ 대치앨리영어"]
                st.code("\n".join(pp))
            
            if st.button("📨 반 전체 발송", type="primary", use_container_width=True):
                success_count = 0
                fail_count = 0
                progress_bar = st.progress(0)
                status_text = st.empty()
                for i, stu in enumerate(students_in_class):
                    if stu['parent_phone']:
                        status_text.text(f"발송 중: {stu['name']} ({i+1}/{len(students_in_class)})")
                        week_att_b = get_week_attendance(stu['id'])
                        msg_parts = [f"안녕하세요, {stu['name']}학부모님.", "자신감을 키우는 대치앨리영어교육입니다.", "", "📅 이번 주 출석 현황"]
                        if week_att_b:
                            for a in week_att_b:
                                msg_parts.append(f"• {a['date']} ({a['day_name']}): {a['status']}")
                        msg_parts += ["", "📝 시험결과"]
                        if class_areas_batch:
                            for a in class_areas_batch:
                                msg_parts.append(f"• {a['area_name']}: -")
                        msg_parts += ["", "❤️ 대치앨리영어"]
                        msg = "\n".join(msg_parts)
                        if batch_template_id:
                            att_str = " | ".join([f"{a['date']} ({a['day_name']}): {a['status']}" for a in week_att_b]) if week_att_b else ""
                            test_str = " | ".join([f"{a['area_name']}: -" for a in class_areas_batch]) if class_areas_batch else ""
                            variables = {
                                "#{학생이름}": stu['name'], "#{출석날짜}": att_str, "#{출석상태}": "",
                                "#{잔여횟수}": "", "#{시험주차}": "", "#{정답률}": test_str,
                                "#{취약영역}": "", "#{리포트월}": "", "#{듣기점수}": "", "#{어휘점수}": "",
                                "#{문법점수}": "", "#{독해점수}": "", "#{AI총평}": "",
                            }
                            res = send_comprehensive_alimtalk(api_key, api_secret, from_number, stu['parent_phone'], batch_pf_id, batch_template_id, variables)
                        else:
                            res = send_message(api_key, api_secret, stu['parent_phone'], from_number, msg)
                        if res.get("success"): success_count += 1
                        else: fail_count += 1
                    progress_bar.progress((i + 1) / len(students_in_class))
                status_text.text(f"✅ 완료! 성공: {success_count}건, 실패: {fail_count}건")
                st.balloons()
        else:
            st.info("등록된 반이 없습니다.")
    
    with t3:
        st.markdown("### 📝 알림톡 템플릿 미리보기")
        st.markdown("""
        <div style="background: #F5F5F5; padding: 20px; border-radius: 20px; max-width: 400px; margin: 0 auto;">
            <div style="background: white; padding: 16px; border-radius: 16px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 12px;">
                    <span style="font-size: 2rem;">🎀</span>
                    <div>
                        <div style="font-weight: 700; color: #D81B60;">대치앨리영어</div>
                        <div style="font-size: 0.75rem; color: #888;">알림톡</div>
                    </div>
                </div>
                <div style="border-top: 1px solid #eee; padding-top: 12px;">
                    <p>안녕하세요, <b>#{학생이름}</b>학부모님.</p>
                    <p>자신감을 키우는 대치앨리영어교육입니다.</p>
                    <br>
                    <p>📅 <b>이번 주 출석 현황</b></p>
                    <p>#{출석날짜}</p>
                    <br>
                    <p>📝 <b>시험결과</b></p>
                    <p>#{정답률}</p>
                    <br>
                    <p style="text-align: right; font-size: 0.8rem; color: #888;">❤️ 대치앨리영어</p>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.info("""
        📌 **알림톡 사용 조건**
        1. 솔라피 카카오 비즈니스 채널 등록 → pfId 발급
        2. 카카오 알림톡 템플릿 승인 → templateId 발급  
        3. 템플릿 미승인 상태면 **일반 문자(LMS)**로 자동 대체 발송됩니다.
        """)

# ─── 푸터 ───────────────────────────────
st.markdown("<br><br><center style='color:#AD1457; font-size:0.8rem;'>© 2025 대치앨리영어 | Made with ❤️ for Allie</center>", unsafe_allow_html=True)